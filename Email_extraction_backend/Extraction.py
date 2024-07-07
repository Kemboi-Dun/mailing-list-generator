from fastapi import FastAPI, File, UploadFile
from fastapi.responses import JSONResponse, FileResponse
from docx import Document
import pandas as pd 
import openpyxl
import re

from fastapi.middleware.cors import CORSMiddleware

from Set_hyperlink import add_hyperlink

app = FastAPI()

origins = [
    "http://localhost:5173"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins = origins,
    allow_credentials = True,
    allow_methods = ["*"],
    allow_headers = ["*"],
)

#Method to extract emails
def extract_emails_from_docx(docx_file):
    doc = Document(docx_file)
    emails = set()
    
    # Extract emails from paragraphs
    for para in doc.paragraphs:
        emails.update(re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',para.text.strip()))
        
    #Extract emails from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                emails.update(re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',cell.text.strip()))
                
    return list(emails)

@app.post("/uploadfile/")
async def create_upload_file(file: UploadFile = File(...)):
    emails = extract_emails_from_docx(file.file)
    return JSONResponse(content={"Emails": emails})

@app.post("/save_excel/")
async def save_file_as_excel(emails:list[str]):
    doc = Document()
    df = pd.DataFrame(emails, columns=["Email"])
    df.to_excel("Excel_Emails.xlsx",  index=False)
    
    file_path = "Excel_Emails.xlsx"
    # doc.save(df)
    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet', filename=file_path)
    

@app.post("/save_word")
async def save_file_as_word(emails: list[str]):
    doc = Document()
    paragraph = doc.add_paragraph()
    for email  in emails:
        add_hyperlink(paragraph, f"mailto:{email}", email)
        paragraph.add_run('; ') # add a separator between emails
    
    file_path = "Word_Emails.docx"
    doc.save(file_path)
    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=file_path)
    # return JSONResponse(content={"Message": "File saved successfully"})


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)